#!/usr/bin/env python3
"""
Parser за екстракција на студентски проекти од docx фајлови.
Секој проект се зачувува како посебен JSON објект.
"""

import os
import json
import re
from docx import Document
from dataclasses import dataclass, asdict
from typing import List, Optional


@dataclass
class ProfessorEvaluation:
    """Оценка од професорот"""
    izmama: str  # Измама - текстуален коментар
    kviz: float  # Квиз поени
    tocnost: float  # Точност фактор
    metap: float  # Метаподатоци фактор
    vkupno: float  # Вкупно = metap * tocnost * kviz
    komentar: str  # Коментар од професорот


@dataclass
class StudentProject:
    """Студентски проект"""
    project_id: str  # Уникатен ID
    project_name: str  # Име на проектот (Тема X: Наслов)
    project_type: str  # noAI, onlyAI, hybrid
    source_file: str  # Од кој фајл е извлечен
    content: str  # Целосна содржина на проектот
    word_count: int  # Број на зборови
    source_count: int  # Број на извори/референци
    professor_eval: Optional[ProfessorEvaluation]  # Оценка од професор


def extract_table_data(table) -> dict:
    """Екстракција на податоци од табела со оценки"""
    data = {}
    if len(table.rows) >= 2:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        values = [cell.text.strip() for cell in table.rows[1].cells]

        for h, v in zip(headers, values):
            if 'измама' in h:
                data['izmama'] = v
            elif 'квиз' in h:
                try:
                    data['kviz'] = float(v.replace(',', '.'))
                except:
                    data['kviz'] = 0.0
            elif 'точност' in h:
                try:
                    data['tocnost'] = float(v.replace(',', '.'))
                except:
                    data['tocnost'] = 1.0
            elif 'метап' in h:
                try:
                    data['metap'] = float(v.replace(',', '.'))
                except:
                    data['metap'] = 1.0
            elif 'вкупно' in h:
                try:
                    data['vkupno'] = float(v.replace(',', '.'))
                except:
                    data['vkupno'] = 0.0
    return data


def count_sources(text: str) -> int:
    """Брои број на референци/извори во текстот"""
    # Барај URL-ови
    urls = re.findall(r'https?://[^\s\]]+', text)
    # Барај нумерирани референци [1], [2], итн.
    numbered_refs = re.findall(r'\[\d+\]', text)
    # Барај bullet points со http
    bullets_with_url = len(re.findall(r'[•\-]\s*http', text))

    # Врати максимум од различните методи
    unique_urls = len(set(urls))
    unique_refs = len(set(numbered_refs))

    return max(unique_urls, unique_refs, bullets_with_url)


def parse_docx(filepath: str, project_type: str) -> List[StudentProject]:
    """Парсирај docx фајл и извлечи ги сите проекти"""
    doc = Document(filepath)
    projects = []

    # Најди ги сите позиции на "Тема X:"
    tema_positions = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^Тема\s+\d+:', text):
            tema_positions.append((i, text))

    # Екстракција на табели
    tables = list(doc.tables)
    table_idx = 0

    filename = os.path.basename(filepath)

    for idx, (start_pos, tema_title) in enumerate(tema_positions):
        # Крај на проектот е почетокот на следниот или крај на документот
        if idx + 1 < len(tema_positions):
            end_pos = tema_positions[idx + 1][0]
        else:
            end_pos = len(doc.paragraphs)

        # Собери ја содржината
        content_parts = []
        for i in range(start_pos, end_pos):
            text = doc.paragraphs[i].text.strip()
            if text and not text.startswith('Оценка:'):
                content_parts.append(text)

        content = '\n\n'.join(content_parts)

        # Број на зборови
        word_count = len(content.split())

        # Број на извори
        source_count = count_sources(content)

        # Оценка од професор (ако има табела)
        professor_eval = None
        if table_idx < len(tables):
            table_data = extract_table_data(tables[table_idx])
            if table_data:
                # Пронајди коментар (обично веднаш по табелата)
                komentar = ""
                for i in range(start_pos, end_pos):
                    text = doc.paragraphs[i].text.strip()
                    if text.startswith('Коментар:'):
                        komentar = text.replace('Коментар:', '').strip()
                        break

                professor_eval = ProfessorEvaluation(
                    izmama=table_data.get('izmama', 'нема'),
                    kviz=table_data.get('kviz', 0.0),
                    tocnost=table_data.get('tocnost', 1.0),
                    metap=table_data.get('metap', 1.0),
                    vkupno=table_data.get('vkupno', 0.0),
                    komentar=komentar
                )
            table_idx += 1

        # Креирај проект
        project_id = f"{project_type}_{idx+1:02d}"
        project = StudentProject(
            project_id=project_id,
            project_name=tema_title,
            project_type=project_type,
            source_file=filename,
            content=content,
            word_count=word_count,
            source_count=source_count,
            professor_eval=professor_eval
        )
        projects.append(project)

    return projects


def main():
    """Главна функција"""
    projects_dir = '/home/artorias/Desktop/OnTP/projects/'
    output_dir = '/home/artorias/Desktop/OnTP/parsed_projects/'

    os.makedirs(output_dir, exist_ok=True)

    # Мапирање на фајлови до типови и суфикси
    file_type_map = {
        'noAI.docx': ('noAI', 'a'),
        'noAI-2.docx': ('noAI', 'b'),
        'onlyAI.docx': ('onlyAI', ''),
        'hybrid.docx': ('hybrid', '')
    }

    all_projects = []
    global_counter = {'noAI': 0, 'onlyAI': 0, 'hybrid': 0}

    for filename, (project_type, suffix) in file_type_map.items():
        filepath = os.path.join(projects_dir, filename)
        if os.path.exists(filepath):
            print(f"Парсирам: {filename}...")
            projects = parse_docx(filepath, project_type)

            # Ажурирај ги ID-овите за да бидат уникатни
            for p in projects:
                global_counter[project_type] += 1
                p.project_id = f"{project_type}_{global_counter[project_type]:02d}{suffix}"

            all_projects.extend(projects)
            print(f"  -> Пронајдени {len(projects)} проекти")

    # Зачувај сите проекти во JSON
    output_file = os.path.join(output_dir, 'all_projects.json')
    projects_data = []
    for p in all_projects:
        data = asdict(p)
        if p.professor_eval:
            data['professor_eval'] = asdict(p.professor_eval)
        projects_data.append(data)

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(projects_data, f, ensure_ascii=False, indent=2)

    print(f"\nВкупно парсирани проекти: {len(all_projects)}")
    print(f"Зачувано во: {output_file}")

    # Статистика по тип
    type_stats = {}
    for p in all_projects:
        type_stats[p.project_type] = type_stats.get(p.project_type, 0) + 1

    print("\nСтатистика по тип:")
    for t, count in type_stats.items():
        print(f"  {t}: {count} проекти")

    # Зачувај секој проект посебно за полесна евалуација
    for p in all_projects:
        project_file = os.path.join(output_dir, f'{p.project_id}.json')
        data = asdict(p)
        if p.professor_eval:
            data['professor_eval'] = asdict(p.professor_eval)
        with open(project_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"\nПоединечни проекти зачувани во: {output_dir}")


if __name__ == '__main__':
    main()
