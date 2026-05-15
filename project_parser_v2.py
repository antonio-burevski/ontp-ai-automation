#!/usr/bin/env python3
"""
Parser v2 за екстракција на студентски проекти од docx фајлови.
Поддржува две фази: training (со оценки) и evaluation (без оценки).
Очекувани фајлови во секоја фаза: no_AI.docx, only_AI.docx, hybrid.docx
"""

import os
import json
import re
from docx import Document
from dataclasses import dataclass, asdict
from typing import List, Optional

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

FILE_TYPE_MAP = {
    'no_AI.docx': 'noAI',
    'only_AI.docx': 'onlyAI',
    'hybrid.docx': 'hybrid',
}


@dataclass
class ProfessorEvaluation:
    izmama: str
    kviz: float
    tocnost: float
    metap: float
    vkupno: float
    komentar: str


@dataclass
class StudentProject:
    project_id: str
    project_name: str
    project_type: str
    source_file: str
    content: str
    word_count: int
    source_count: int
    professor_eval: Optional[ProfessorEvaluation]


def extract_table_data(table) -> dict:
    data = {}
    if len(table.rows) >= 2:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        values = [cell.text.strip() for cell in table.rows[1].cells]

        for h, v in zip(headers, values):
            if 'измама' in h:
                data['izmama'] = v
            elif 'квиз' in h:
                try:
                    data['kviz'] = float(v.replace(',', '.'))
                except ValueError:
                    data['kviz'] = 0.0
            elif 'точност' in h:
                try:
                    data['tocnost'] = float(v.replace(',', '.'))
                except ValueError:
                    data['tocnost'] = 1.0
            elif 'метап' in h:
                try:
                    data['metap'] = float(v.replace(',', '.'))
                except ValueError:
                    data['metap'] = 1.0
            elif 'вкупно' in h:
                try:
                    data['vkupno'] = float(v.replace(',', '.'))
                except ValueError:
                    data['vkupno'] = 0.0
    return data


def count_sources(text: str) -> int:
    urls = re.findall(r'https?://[^\s\]]+', text)
    numbered_refs = re.findall(r'\[\d+\]', text)
    bullets_with_url = len(re.findall(r'[•\-]\s*http', text))
    return max(len(set(urls)), len(set(numbered_refs)), bullets_with_url)


def parse_docx(filepath: str, project_type: str) -> List[StudentProject]:
    doc = Document(filepath)
    projects = []

    tema_positions = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if re.match(r'^Тема\s+\d+:', text):
            tema_positions.append((i, text))

    tables = list(doc.tables)
    table_idx = 0
    filename = os.path.basename(filepath)

    for idx, (start_pos, tema_title) in enumerate(tema_positions):
        end_pos = tema_positions[idx + 1][0] if idx + 1 < len(tema_positions) else len(doc.paragraphs)

        content_parts = []
        for i in range(start_pos, end_pos):
            text = doc.paragraphs[i].text.strip()
            if text and not text.startswith('Оценка:'):
                content_parts.append(text)

        content = '\n\n'.join(content_parts)
        word_count = len(content.split())
        source_count = count_sources(content)

        professor_eval = None
        if table_idx < len(tables):
            table_data = extract_table_data(tables[table_idx])
            if table_data:
                komentar = ""
                for i in range(start_pos, end_pos):
                    text = doc.paragraphs[i].text.strip()
                    if text.startswith('Коментар:'):
                        komentar = text.replace('Коментар:', '').strip()
                        break

                professor_eval = ProfessorEvaluation(
                    izmama=table_data.get('izmama', 'нема'),
                    kviz=table_data.get('kviz', 0.0),
                    tocnost=table_data.get('tocnost', 1.0),
                    metap=table_data.get('metap', 1.0),
                    vkupno=table_data.get('vkupno', 0.0),
                    komentar=komentar
                )
            table_idx += 1

        project_id = f"{project_type}_{idx + 1:02d}"
        project = StudentProject(
            project_id=project_id,
            project_name=tema_title,
            project_type=project_type,
            source_file=filename,
            content=content,
            word_count=word_count,
            source_count=source_count,
            professor_eval=professor_eval
        )
        projects.append(project)

    return projects


def parse_phase(phase: str) -> List[dict]:
    """Parse all docx files for a given phase (training or evaluation)."""
    phase_dir = os.path.join(BASE_DIR, 'projects', phase)
    output_dir = os.path.join(BASE_DIR, 'parsed_projects_v2', phase)
    os.makedirs(output_dir, exist_ok=True)

    all_projects = []

    for filename, project_type in FILE_TYPE_MAP.items():
        filepath = os.path.join(phase_dir, filename)
        if not os.path.exists(filepath):
            print(f"  Предупредување: {filepath} не постои, прескокнувам.")
            continue

        print(f"  Парсирам: {filename} ({project_type})...")
        projects = parse_docx(filepath, project_type)

        for i, p in enumerate(projects, 1):
            p.project_id = f"{phase}_{project_type}_{i:02d}"

        all_projects.extend(projects)
        print(f"    -> Пронајдени {len(projects)} проекти")

    projects_data = []
    for p in all_projects:
        data = asdict(p)
        if p.professor_eval:
            data['professor_eval'] = asdict(p.professor_eval)
        projects_data.append(data)

        project_file = os.path.join(output_dir, f'{p.project_id}.json')
        with open(project_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    all_file = os.path.join(output_dir, 'all_projects.json')
    with open(all_file, 'w', encoding='utf-8') as f:
        json.dump(projects_data, f, ensure_ascii=False, indent=2)

    print(f"  Вкупно: {len(all_projects)} проекти -> {all_file}")
    return projects_data


def main():
    import argparse
    parser = argparse.ArgumentParser(description='Parser v2 за студентски проекти')
    parser.add_argument('phase', choices=['training', 'evaluation', 'both'],
                        help='Која фаза да се парсира')
    args = parser.parse_args()

    phases = ['training', 'evaluation'] if args.phase == 'both' else [args.phase]
    for phase in phases:
        print(f"\n=== Парсирање: {phase.upper()} ===")
        parse_phase(phase)


if __name__ == '__main__':
    main()
